#用于模拟浏览器
from selenium import webdriver
import unittest, time#延时用
#用于读取docx文档
import docx
from docutils.nodes import paragraph
#beautiful 用于解析网页
import bs4 as bs
import http.client as httplib
#快速请求
import httplib2
import urllib.request
#普通请求
import http.cookiejar as cookielib
import urllib.parse
import json #json数据格式，用于写入数据库链接配置信息
from types import CodeType
#正则匹配文本
import re
from logging import exception
#读写xml文件
try:
  import xml.etree.cElementTree as ET
except ImportError:
  import xml.etree.ElementTree as ET

import xlrd
from xlutils.copy import copy as xlwtcopy

t1 = time.clock()

def requestNeedurl(needurl,postData,headers,CookieOpener):
    needurl = needurl
    headers['Accept'] = 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8' 
    if headers.get('Content-Type') == None:
        pass
    else:
        del headers['Content-Type']
    encodepostData = urllib.parse.urlencode(postData).encode('utf-8')
    req = urllib.request.Request(url = needurl,data = encodepostData,headers = headers)
    response = CookieOpener.open(req)
    return response
def getpermanentID(rooturl,headers,loginurl): 
    req = urllib.request.Request(url = loginurl,headers = headers)
    response = urllib.request.urlopen(req)
    soup = bs.BeautifulSoup(response.read().decode(),"lxml")
    permanentID = soup.form.find(id="permanentId")["value"]
    return permanentID
def loginInSystem(requestInfo):
    permanentID = getpermanentID(rooturl=requestInfo.rooturl,headers=headers,loginurl=requestInfo.loginurl)
    postData = {'account':requestInfo.webname,'password':requestInfo.webpsw,'cultureCode':'zh-CN','permanentId':permanentID}
    print('requestInfo.loginurl:%s'%(requestInfo.loginurl))
    response = requestNeedurl(needurl  = requestInfo.loginurl ,postData = postData,headers = headers,CookieOpener = requestInfo.CookieOpener)
    fullcookies = ''
    for item in cookie:
            singlecookie = item.name + '=' + item.value + ';'
            fullcookies = fullcookies + singlecookie
    requestInfo.headers['Cookie'] = fullcookies        
    return response
def requestNeedurlByJsonPost(needurl,postData,headers,host,port):
    needurl = needurl
    headers['Content-Type'] = 'application/json'
    headers['Accept'] = 'application/json, text/javascript, */*; q=0.01'
    conn = httplib.HTTPConnection(host,port)
    conn.request("POST", needurl, json.JSONEncoder().encode(postData), headers)
    response = conn.getresponse()
    return response



class UrlRequest:
    def __init__(self,rooturl,host,port,webname,webpsw,tmpFilePath,showMode,cookie,headers,loginurl,CookieOpener,http):
        self.rooturl = rooturl
        self.host = host
        self.port = port
        self.webname = webname
        self.webpsw = webpsw
        self.tmpFilePath = tmpFilePath 
        self.showMode = showMode
        self.cookie = cookie
        self.headers = headers
        self.loginurl = loginurl
        self.CookieOpener = CookieOpener  
        self.http = http       
def checkContentInMimeTypes(checkThing,mimetypepath,addSet):
    ''' for check content like xls in mimetype configfile,and delete old config,
        set new config(can download with no windows),return update result and old  content which was overwrite
        checkThing : for xls、xlsx、pdf,be sure as same as addSet config
        mimetypepath:the config file in local for webdirver firefox
        addset:for checkthings ,you need each set for checktings,dict it is 
    '''
    checkState = 0
    oldCheckContent = ''
    fileobject = open(mimetypepath,'rb')
    fileReadContent = fileobject.readlines()
    fileobject.close()
    
    fileobject = open(mimetypepath,'wb')
    checkName = r'urn:mimetype:application/' + checkThing
    for line in fileReadContent:
        lineContent = line.decode('utf-8')              
        if checkState > 0:
            if  r'</RDF:Description>' in lineContent:
                checkState = 0
                oldCheckContent = oldCheckContent + lineContent
            else:
                checkState += 1
                oldCheckContent = oldCheckContent + lineContent
        else:
            if checkName in lineContent:
                checkState = 1
                oldCheckContent = oldCheckContent + lineContent
                if  r'</RDF:Description>' in lineContent:
                    checkState = 0
                else:
                    pass
            else:
                if r'</RDF:RDF>' in lineContent:
                    
                    fileobject.write(addSet.get(checkThing).encode('utf-8')) 
                    fileobject.write(line)
                else:
                    fileobject.write(line)
    fileobject.close()
        
    fileobject = open(mimetypepath,'rb')
    newReadContent = fileobject.read().decode('utf-8')
    fileobject.close()                
    return  newReadContent,oldCheckContent    
def writeMimetype(mimetypepath): 
    ''' use def checkContentInMimeTypes overwrite  mimetype config file one by one
    mimetypepath:the config file in local for webdirver firefox '''   
    mimetypepath = r'C:\Users\Administrator\AppData\Roaming\Mozilla\Firefox\Profiles\hdjophne.default\mimeTypes.rdf'
    addSet = {} 
    addSet['xls'] =  r'''<RDF:Description RDF:about="urn:mimetype:application/xls" NC:value="application/xls" NC:fileExtensions="xls">  </RDF:Description>
        <RDF:Description RDF:about="urn:mimetype:handler:application/xls" NC:saveToDisk="true" NC:alwaysAsk="false" ></RDF:Description>'''
    addSet['xlsx'] = r'''<RDF:Description RDF:about="urn:mimetype:application/xlsx" NC:value="application/xlsx" NC:fileExtensions="xlsx">  </RDF:Description>
        <RDF:Description RDF:about="urn:mimetype:handler:application/xlsx" NC:saveToDisk="true" NC:alwaysAsk="false" ></RDF:Description>'''
    newcheckXls,oldcheckXls = checkContentInMimeTypes(checkThing='xls',mimetypepath=mimetypepath,addSet=addSet)
    #print('newcheckXls:%s'%(newcheckXls))
    #print('oldcheckXls:%s'%(oldcheckXls))
    newcheckXlsx,oldcheckXlsx = checkContentInMimeTypes(checkThing='xlsx',mimetypepath=mimetypepath,addSet=addSet)
    #print('newcheckXlsx:%s'%(newcheckXlsx))
    #print('oldcheckXlsx:%s'%(oldcheckXlsx))
def createDriver(requestInfo):
    '''create firefox webdriver and phantomjs webdriver 
        showMode:['show','hide']
    '''
    profile_directory = r'C:\Users\Administrator\AppData\Roaming\Mozilla\Firefox\Profiles\hdjophne.default'
    writeMimetype(mimetypepath=profile_directory)
    profile = webdriver.FirefoxProfile(profile_directory)
    profile.set_preference('browser.download.dir', r'E:\KM\GITPROJECT\HumanResourcesAss\AssForGaia\src\Guru\tmpfile')
    if requestInfo.showMode == 'show':    
        driver = webdriver.Firefox(profile)
    elif requestInfo.showMode == 'hide':
        driver = webdriver.PhantomJS('phantomjs')
    loginurl = requestInfo.rooturl + r'/Account/Logon?'
    driver.get(loginurl)
    time.sleep(2)
    driver.find_element_by_id('account_INPUT_IDENTITY').send_keys(requestInfo.webname)
    driver.find_element_by_id('password').send_keys(requestInfo.webpsw)
    driver.find_element_by_class_name('btn-lg').click()
    time.sleep(2)
    return driver
def readPublicByDoc(requestInfo,fileName):
    document = docx.Document(fileName)
    publicContent = []
    for table in document.tables:
        tableContent = {}
        codeValue = {}
        if table.cell(0,0).text == '公用代码配置':
            tableValue = []
            for row in table.rows:
                rowContent = []
                for cell in row.cells:
                    rowContent.append(cell.text)
                if rowContent[0] in ['公用代码配置','代码所属表','代码编号']:
                    pass
                elif rowContent[0] == '代码类型':
                    tableContent['codeType'] = rowContent[1]
                elif rowContent[0] == '代码名称':
                    tableContent['codeName'] = rowContent[1]
                else:
                    codeValue['codeItemID'] = rowContent[0]
                    codeValue['codeItemValue'] = rowContent[1]
                    tableValue.append(codeValue.copy())
            tableContent['codeItem'] = tableValue.copy()
            publicContent.append(tableContent.copy())
    #print('publicContent:%s'%(publicContent))
    return  publicContent     
def readExtendByDoc(requestInfo,fileName):
    document = docx.Document(fileName)
    extendContent = []
    for table in document.tables:
        tableContent = {}
        codeItem = []
        if table.cell(0,0).text == '业务数据配置':
            tableValue = []
            tableValueName = []
            for row in table.rows:
                rowContent = []
                for cell in row.cells:
                    rowContent.append(cell.text.replace('\u3000',''))
                if rowContent[1] in ['业务数据配置','项目']:
                    pass
                elif rowContent[1] == '业务数据类型':
                    tableContent['extendType'] = rowContent[2]
                elif rowContent[1] == '群组名称':
                    tableContent['extendName'] = rowContent[2]
                elif rowContent[1] == '重复限制':
                    tableContent['repeatLimit'] = rowContent[2]                    
                elif rowContent[1] == '群组描述':
                    tableContent['extendDesc'] = rowContent[2]
                elif rowContent[1] == '列字段名称':                                          
                    tableValueName = rowContent
                else:
                    tableValue.append(rowContent.copy())
            for singleValue in  tableValue:
                singleCodeItem = {}
                i = 0
                for i in range(len(singleValue)):
                    singleCodeItem[tableValueName[i]] = singleValue[i]   
                codeItem.append(singleCodeItem.copy())    
            tableContent['codeItem'] = codeItem.copy()
            extendContent.append(tableContent.copy())
    #print('extendContent:%s'%(extendContent))
    return  extendContent  
def queryPublicGroup(requestInfo):
    listPublicGroupName = []
    needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/PayrollPublicCodeMaintain.aspx'    
    content = requestInfo.http.request(uri = needurl,method = 'GET', headers=requestInfo.headers)[1]
    readSoup = bs.BeautifulSoup(content.decode(),'lxml')
    listPublicGroupName = listPublicGroupName + [singleGroup.find_all('span')[2].string for singleGroup in readSoup.find('div',id='M_UWebTreePublicCode_1').find_all('div') if len(singleGroup.find_all('span')[2].string)>0].copy()    
    needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/PayrollPublicCodeMaintain.aspx?Module=INS'     
    content = requestInfo.http.request(uri = needurl,method = 'GET', headers=requestInfo.headers)[1]
    readSoup = bs.BeautifulSoup(content.decode(),'lxml')
    listPublicGroupName = listPublicGroupName + [singleGroup.find_all('span')[2].string for singleGroup in readSoup.find('div',id='M_UWebTreePublicCode_1').find_all('div') if len(singleGroup.find_all('span')[2].string)>0].copy()    
    #print('listPublicGroupName:%s'%(listPublicGroupName))
    return listPublicGroupName
def addPublicGroup(driver,requestInfo,codeType,addPublicGroupName,listPublicGroupName):
    if addPublicGroupName in listPublicGroupName:
        print('Group name %s in PublicGroupName list,Add ignore!'%(addPublicGroupName))
    else:
        if codeType == '薪资公用代码':
            codeTypeValue ='PAY'
        elif codeType == '保险共用代码':
            codeTypeValue = 'INS'
        if  codeTypeValue is not None and len(addPublicGroupName) > 0:   
            needurl = requestInfo.rooturl +r'/ePayroll/PayrollParameterSetting/PayrollPublicCodeMaintainEdit.aspx?Module=' + codeTypeValue
            driver.get(needurl)
            time.sleep(2)
            driver.find_element_by_id('txtCode').send_keys(addPublicGroupName)
            time.sleep(2)
            driver.find_element_by_id('cmdAdd').click()
            time.sleep(2)
            message = driver.switch_to_alert().text
            if message == '新增成功！':
                print('Group Add Success:%s'%(addPublicGroupName))
                driver.switch_to_alert().accept()
            else:
                print('Group Add Failed:%s,Because of:%s'%(addPublicGroupName,message))
                driver.switch_to_alert().accept()
        else:
            print('codeTypeValue is None or addPublicGroupName:%s is null,Add ignore!'%(addPublicGroupName))
def updatePublicGroupByDoc(driver,requestInfo,listPublic):
    listPublicGroupName = queryPublicGroup(requestInfo)
    for singlePublic in listPublic:
        addPublicGroupCodeType = singlePublic.get('codeType')
        addPublicGroupName = singlePublic.get('codeName')
        if addPublicGroupCodeType in ['保险共用代码','薪资公用代码']:
            addPublicGroup(driver,requestInfo,addPublicGroupCodeType,addPublicGroupName,listPublicGroupName)
        else:
            print('for group:%s  Code Type is not 保险共用代码,薪资公用代码 ,add ignore'%(addPublicGroupName))
def queryExtendGroup(requestInfo):
    listExtendGroup = []
    needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/ExtendDataGroupSetting.aspx'    
    content = requestInfo.http.request(uri = needurl,method = 'GET', headers=requestInfo.headers)[1]
    readSoup = bs.BeautifulSoup(content.decode(),'lxml')
    singletr = readSoup.find('thead',class_='ig_e531842b_r1 WebGridText ig_e531842b_r4').find('tr')
    extendNamelist = [singletd.find('nobr').string for singletd in singletr.find_all('th')].copy()
    extendValuelist = []
    for singletr in readSoup.find('tbody',class_='ig_e531842b_r1 WebGridText ig_e531842b_r4').find_all('tr'):
        extendValuelist.append([singletd.find('nobr').string for singletd in singletr.find_all('td')].copy())
    if  len(extendValuelist) > 0: 
        for singleExtendValue in extendValuelist:
            singleExtendGroup = {}
            for i in range(len(singleExtendValue)):
                if extendNamelist[i] is not None and len(extendNamelist[i]) > 0 and extendNamelist[i] in ['群组名称','业务数据群组类型','群组描述','群组编号']:
                    singleExtendGroup[extendNamelist[i]] = singleExtendValue[i]
            listExtendGroup.append(singleExtendGroup.copy())
    else:
        print('No extend group list')      
    print('listExtendGroup:%s'%(listExtendGroup))
    return listExtendGroup
def queryExtendDetail(requestInfo,groupID):
    listExtendDetail = [] 
    needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/ExtendDataGroupSettingDetail.aspx?GroupID=' +  groupID   
    content = requestInfo.http.request(uri = needurl,method = 'GET', headers=requestInfo.headers)[1]
    readSoup = bs.BeautifulSoup(content.decode(),'lxml')
    singletr = readSoup.find('thead',class_='ig_e531842b_r1 WebGridText ig_e531842b_r4').find('tr')
    extendDetailNamelist = [singletd.find('nobr').string for singletd in singletr.find_all('th')].copy()
    extendDetailValuelist = []
    for singletr in readSoup.find('tbody',class_='ig_e531842b_r1 WebGridText ig_e531842b_r4').find_all('tr'):
        extendDetailValuelist.append([singletd.find('nobr').string for singletd in singletr.find_all('td')].copy())
    if  len(extendDetailValuelist) > 0: 
        for singleExtendDetailValue in extendDetailValuelist:
            singleExtendDetail = {}
            for i in range(len(singleExtendDetailValue)):
                if extendDetailNamelist[i] is not None and len(extendDetailNamelist[i]) > 0 and extendDetailNamelist[i] in ['序号','字段名称','数据类型','是否必填','显示方式','数据关联','字段编号']:
                    singleExtendDetail[extendDetailNamelist[i]] = singleExtendDetailValue[i].replace('\xa0','')
            listExtendDetail.append(singleExtendDetail.copy())
    else:
        print('No extend detail list')      
    print('listExtendDetail:%s'%(listExtendDetail))
    return listExtendDetail

              
def addExtendGroup(driver,requestInfo,addSingleExtendGroup,listExtendGroup):  
    addExtendGroupName = addSingleExtendGroup.get('extendName')
    listExtendGroupName = [group.get('群组名称') for group in listExtendGroup]
    if addExtendGroupName in listExtendGroupName:
        print('ExtendGroupName %s in ExtendGroupName list,Add ignore!'%(addExtendGroupName))
    elif len(addExtendGroupName) > 0:
        needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/ExtendDataGroupEdit.aspx'
        driver.get(needurl)
        time.sleep(2)
        driver.find_element_by_id('txtGroupName').send_keys(addExtendGroupName)
        addExtendGroupDesc = addSingleExtendGroup.get('extendDesc')
        driver.find_element_by_id('txtGroupDescription').send_keys(addExtendGroupDesc)
        addExtendGroupRepeatLimit = addSingleExtendGroup.get('repeatLimit')
        if addExtendGroupRepeatLimit in ['无限制','按日','按月','按季度','按年']:
            driver.find_element_by_xpath("//select[@id='DropDownListLimitType']/option[@title='"+ addExtendGroupRepeatLimit +"']").click()
        else:
            print('Group repeat limit:%s is not variable!'%(addExtendGroupRepeatLimit))
        addExtendGroupType = addSingleExtendGroup.get('extendType') 
        if addExtendGroupType == '组织业务数据':
            driver.find_element_by_id('rlType_0').click()
        elif addExtendGroupType == '员工业务数据':
            driver.find_element_by_id('rlType_1').click()
        else:
            print('Group type :%s is not variable!'%(addExtendGroupType))
        time.sleep(2)
        driver.find_element_by_id('btnAdd').click()
        time.sleep(2)
        message = driver.switch_to_alert().text
        print(message)
        if message == '新增成功！':
            print('Group Add Success:%s'%(addExtendGroupName))
            driver.switch_to_alert().accept()
        else:
            print('Group Add Failed:%s,Because of:%s'%(addExtendGroupName,message))
            driver.switch_to_alert().accept()
    else:
        print('addExtendGroupName:%s is null,Add ignore!'%(addExtendGroupName))
        
def addExtendDetail(driver,requestInfo,addSingleExtendDetail,listExtendDetail,groupID):        
    addExtendDetailName = addSingleExtendDetail.get('列字段名称')
    listExtendDetailName = [group.get('字段名称') for group in listExtendDetail]
    if addExtendDetailName in listExtendDetailName:
        print('ExtendDetailName %s in ExtendDetailName list,Add ignore!'%(addExtendDetailName))
    elif len(addExtendDetailName) > 0:
        needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/ExtendDataGroupDetailEdit.aspx?GroupID=' + groupID
        driver.get(needurl)
        time.sleep(2)
        driver.find_element_by_id('txtColumnName').send_keys(addExtendDetailName)
        
        addExtendDetailOrder = addSingleExtendDetail.get('序号')
        driver.find_element_by_id('txtOrder').send_keys(addExtendDetailOrder)
        
        addExtendDetailShowMode = addSingleExtendDetail.get('显示方式') 
        if addExtendDetailShowMode == '文本框':
            driver.find_element_by_id('rdoShowMode_0').click()
        elif addExtendDetailShowMode == '选择框':
            driver.find_element_by_id('rdoShowMode_1').click()
        else:
            print('GroupDetail show mode :%s is not variable!'%(addExtendDetailShowMode))     

        addExtendDetailMustFill = addSingleExtendDetail.get('是否必填') 
        if addExtendDetailMustFill == '是':
            driver.find_element_by_id('rdoMustFill_0').click()
        elif addExtendDetailMustFill == '否':
            driver.find_element_by_id('rdoMustFill_1').click()
        else:
            print('GroupDetail MustFill :%s is not variable!'%(addExtendDetailMustFill)) 

        addExtendDetailColumnType = addSingleExtendDetail.get('数据类型') 
        if addExtendDetailColumnType == '文本':
            driver.find_element_by_id('rdoColumnType_0').click()
        elif addExtendDetailColumnType == '日期':
            driver.find_element_by_id('rdoColumnType_1').click()
        elif addExtendDetailColumnType == '数字':
            driver.find_element_by_id('rdoColumnType_2').click()
        else:
            print('GroupDetail ColumnType :%s is not variable!'%(addExtendDetailColumnType)) 
            
        addExtendDetailPublicCode = addSingleExtendDetail.get('数据关联') 
        if  addExtendDetailShowMode == '选择框' and addExtendDetailPublicCode is not None and len(addExtendDetailPublicCode) > 0: 
            driver.find_element_by_xpath("//select[@id='drpPublicCode']/option[@title='"+ addExtendDetailPublicCode +"']").click()
    
        driver.find_element_by_id('btnAdd').click()
        time.sleep(2)
        message = driver.switch_to_alert().text
        print(message)
        if message == '新增成功！':
            print('GroupDetail Add Success:%s'%(addExtendDetailName))
            driver.switch_to_alert().accept()
        else:
            print('GroupDetail Add Failed:%s,Because of:%s'%(addExtendDetailName,message))
            driver.switch_to_alert().accept()
    else:
        print('addExtendDetailName:%s is null,Add ignore!'%(addExtendDetailName))        
           
def importPublicCodeDetail(driver,requestInfo,listExtendByDoc,listPublicByDoc):
    listExtendGroup = queryExtendGroup(requestInfo)
    pc_ExtendGroup = {'repeatLimit': '无限制', 'extendName': '更新公用代码明细', 'extendDesc': 'Sys config for update publiccodeitem', 'extendType': '组织业务数据', 'codeItem': [{'列字段名称': '公用代码类型', '显示方式': '文本框', '数据类型': '文本', '是否必填': '否', '数据关联': '', '序号': '1'}, {'列字段名称': '公用代码编号', '显示方式': '文本框', '数据类型': '文本', '是否必填': '否', '数据关联': '', '序号': '2'},{'列字段名称': '公用代码名称', '显示方式': '文本框', '数据类型': '文本', '是否必填': '否', '数据关联': '', '序号': '3'}]}
    #firt step:add extend group
    addExtendGroup(driver,requestInfo,pc_ExtendGroup,listExtendGroup)
    #second step:add extend detail
    listExtendGroup = queryExtendGroup(requestInfo)
    pc_SysExtendGroup = [singleExtendGroup for singleExtendGroup in listExtendGroup if singleExtendGroup.get('群组名称')==pc_ExtendGroup.get('extendName')]
    pc_ExtendDetail = pc_ExtendGroup.get('codeItem')
    pc_SysExtendGroup_GroupID = pc_SysExtendGroup[0].get('群组编号')
    pc_SysExtendGroup_ExtendDetail = queryExtendDetail(requestInfo,pc_SysExtendGroup_GroupID)
    if pc_ExtendDetail is not None and len(pc_ExtendDetail) > 0:
        for single_pc_ExtendDetail in pc_ExtendDetail:
            addExtendDetail(driver,requestInfo,single_pc_ExtendDetail,pc_SysExtendGroup_ExtendDetail,pc_SysExtendGroup_GroupID)    
    #third step:download template
    needurl = requestInfo.rooturl + r'/ePayroll/PersonalPayrollInformationManage/OrgBusinessDataImport.aspx'
    driver.get(needurl)
    time.sleep(2)
    driver.find_element_by_xpath("//select[@id='drpGroup']/option[@title='"+ pc_ExtendGroup.get('extendName') +"']").click()
    driver.find_element_by_id('btnQuery').click()
    time.sleep(2)
    driver.find_element_by_xpath("//select[@class='pager-sizes']/option[@title='200']").click()
    time.sleep(2)
    tmpPageSource = driver.page_source    
    readSoup = bs.BeautifulSoup(tmpPageSource,"html.parser")
    recordTotalNum =readSoup.find(id = 'grdNavigator_lblRecordsCount').string
    totalPagesNum = int(int(recordTotalNum)/200)
    for i in range(totalPagesNum):
        driver.find_element_by_id('btnQuery').click()
        time.sleep(2)
        driver.find_element_by_id('chkAll').click()
        time.sleep(2)
        driver.find_element_by_id('btnDelete').click()
        time.sleep(2)
        msg = driver.switch_to_alert().text
        print(msg)
        driver.switch_to_alert().accept()
    driver.find_element_by_id('downTemplate').click()   
    time.sleep(2)
    now_handle = driver.current_window_handle
    all_handles = driver.window_handles
    for handle in all_handles:
        if handle in driver.window_handles:
            driver.switch_to_window(handle)
            if driver.current_window_handle == now_handle:
                pass
            else:
                driver.close()     
        else:
            pass
    driver.switch_to_window(now_handle)
    tmpPageSource = driver.page_source    
    readSoup = bs.BeautifulSoup(tmpPageSource,"html.parser")
    downLoadUrl = requestInfo.rooturl + re.findall(r"/CommonPage.*\.xls",str(readSoup))[0]
    response,content = requestInfo.http.request(uri = downLoadUrl,method = 'POST', headers=requestInfo.headers)
    readSoup = bs.BeautifulSoup(content.decode('utf-8'),"html.parser")
    trueLoadUrl = requestInfo.rooturl + re.findall(r"/DownLoadDir.*\.xls",str(readSoup))[0]
    #print(content.decode('utf-8'))
    filename = requestInfo.tmpFilePath + r'/ImportExtendDataForPublicCodeUpdate.xml'
    newfilename = requestInfo.tmpFilePath + r'\NewdataExtendDataForPublicCodeUpdate.xml'
    urllib.request.urlretrieve(url=trueLoadUrl, filename=filename)
    
    #for singleExtend in listExtendByDoc:
    row_former = '<ss:Row>\r\n' 
    col_former = '<ss:Cell>\r\n<ss:Data ss:Type="String">'
    row_end = '</ss:Row>\r\n'
    col_end = '</ss:Data>\r\n</ss:Cell>\r\n'                  
    defaultOrgCode = '01'
    defaultEffectDate = '2015-01-01'
    addContent = ''
    i = 1 #从第二行开始
    for sinPC in listPublicByDoc:
        print('sinPC:%s'%(sinPC))
        PCType = sinPC.get('codeType')
        PCName = sinPC.get('codeName')
        listCodeItem = sinPC.get('codeItem') 
        if PCType in ['薪资公用代码','保险共用代码']:
            for sinItem in listCodeItem:
                print('sinItem:%s'%(sinItem))
                itemID = sinItem.get('codeItemID')
                itemValue = sinItem.get('codeItemValue')
                addContent = addContent + row_former + col_former + col_end   
                addContent = addContent + col_former + defaultOrgCode + col_end 
                addContent = addContent + col_former + defaultEffectDate + col_end 
                addContent = addContent + col_former + PCName + col_end 
                addContent = addContent + col_former + itemID + col_end 
                addContent = addContent + col_former + itemValue + col_end + row_end
    #print('addContent:%s'%(addContent))
     
    fileobject = open(filename,'rb')
    fileReadContent = fileobject.readlines()
    fileobject.close()
    sheetState = 0
    tableState = 0
    fileobject = open(newfilename,'wb')
    for line in fileReadContent:
        lineContent = line.decode('utf-8')
        #print('lineContent:%s'%(lineContent))         
        if  sheetState == 0:      
            if  r'Worksheet' in lineContent and r'OrgBusinessDataImport' in lineContent:
                sheetState = 1
        elif sheetState > 0:
            if  r'<ss:Table>' in lineContent:
                tableState = 1 
            if  r'</ss:Worksheet>' in lineContent:
                sheetState = 0
            if tableState > 0 :
                if  r'</ss:Table>' in lineContent:
                    tableState = 0
                    line = addContent.encode('utf-8') + line
                    
                else:
                    pass
        #print('sheetState:%s'%(sheetState)) 
        #print('tableState:%s'%(tableState))         
        #print('line:%s'%(line))        
        fileobject.write(line)            
    fileobject.close()            

    
    needurl = requestInfo.rooturl + r'/ePayroll/PersonalPayrollInformationManage/OrgBusinessDataImportAction.aspx?GroupId=' + pc_SysExtendGroup_GroupID
    driver.get(needurl)
    time.sleep(2)
    driver.find_element_by_id('lblimportdata').click()
    time.sleep(2)
    driver.find_element_by_id('ucUploadFile_myFile').send_keys(newfilename)
    driver.find_element_by_id('ucUploadFile_btnImput').click()
    time.sleep(2)
    driver.find_element_by_id('btnSave').click()
    
    
    #msg = driver.switch_to_alert().accept()
    #print('msg:%s'%(msg))   
    
def updateExtendGroup(driver,requestInfo,listExtendByDoc):
    listExtendGroup = queryExtendGroup(requestInfo)
    for sinExtend in listExtendByDoc:
        addExtendGroup(driver,requestInfo,sinExtend,listExtendGroup)
    listExtendGroup = queryExtendGroup(requestInfo)
    for cur_ExtendGroup in listExtendByDoc:
        cur_SimExtendGroup = [singleExtendGroup for singleExtendGroup in listExtendGroup if singleExtendGroup.get('群组名称')==cur_ExtendGroup.get('extendName')]
        cur_ExtendDetail = cur_ExtendGroup.get('codeItem')
        print('cur_ExtendGroup:%s'%(cur_ExtendGroup))  
        cur_ExtendGroup_GroupID = cur_SimExtendGroup[0].get('群组编号')
        cur_ExtendGroup_ExtendDetail = queryExtendDetail(requestInfo,cur_ExtendGroup_GroupID)
        if cur_ExtendDetail is not None and len(cur_ExtendDetail) > 0:
            for sin_cur_ExtendDetail in cur_ExtendDetail:
                addExtendDetail(driver,requestInfo,sin_cur_ExtendDetail,cur_ExtendGroup_ExtendDetail,cur_ExtendGroup_GroupID) 
def queryFormula(driver,requestInfo,CFGfilename):
    
    workbook = xlrd.open_workbook(CFGfilename)
    ws = workbook.sheet_by_name('OPT010薪资函数')
    rowlen,collen = ws.usp_get_len()
    tablename = []
    tablecontent = []
    for i in range(rowlen):
        rowcontent = []
        for j in range(collen):
            if i == 0:
                cellcontent = ws.cell(i,j).value
                tablename.append(ws.cell(i,j).value )
            else:
                cellcontent = ws.cell(i,j).value
                rowcontent.append(ws.cell(i,j).value)
        if i != 0:        
            tablecontent.append(rowcontent.copy())
    i = 0
    tableFormu = []
    dictFormu = {}
    for sinFor in tablecontent:
        #print('sinFor:%s'%(sinFor))
        for i in range(len(tablename)):
            dictFormu[tablename[i]] = sinFor[i]
        tableFormu.append(dictFormu.copy())
    print('tableFormu:%s'%(tableFormu))
    return tableFormu

def updateFormula(driver,requestInfo,listFormuByDoc):       
    needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/PayrollScriptEdit.aspx?optype=add&addtype=FUNCTIONASSISTANTKEY'
    driver.get(needurl)   
    for sinForm in listFormuByDoc:
        formu_name = sinForm.get('名称')
        formu_usetype = sinForm.get('使用状态')
        formu_type = sinForm.get('类型')
        formu_enname = sinForm.get('调用代码')
        formu_code = sinForm.get('函数内容')
        formu_help = sinForm.get('帮助')            
        if formu_usetype == '使用中' and formu_type == '辅助函数':
            driver.get(needurl) 
            time.sleep(2)
            driver.find_element_by_id('txtName').send_keys(formu_name)
            driver.find_element_by_id('txtInnerName').send_keys(formu_enname)
            driver.find_element_by_id('txtContent').send_keys(formu_code)
            driver.find_element_by_id('txtHelper').send_keys(formu_help)
            driver.find_element_by_id('btnSave').click()
            time.sleep(2)
            try:
                msg = driver.switch_to_alert().text
            except(Exception):
                print('no msg formu_name:%s'%(formu_name))
            else:
                if msg == '新增成功':
                    print('Add Success formu_name：%s'%(formu_name))
                    driver.switch_to_alert().accept()
                else:
                    print('Add Failed formu_name:%s;Because of:%s'%(formu_name,msg))
                    driver.switch_to_alert().accept()
        else:
            print('other type formu_name%s'%(formu_name)) 
    
def addInsArea(driver,requestInfo,sinInsArea):
    needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/InsureSystemSetEdit.aspx'
    driver.get(needurl)
    time.sleep(2)
    driver.find_element_by_id('txtAreaName').send_keys(sinInsArea)
    driver.find_element_by_id('btnAdd').click()
    time.sleep(2)    
    try:
        msg = driver.switch_to_alert().text
    except(Exception):
        print('no msg sinInsArea:%s'%(sinInsArea))
    else:
        if msg == '新增成功':
            print('Add Success sinInsArea：%s'%(sinInsArea))
            driver.switch_to_alert().accept()
        else:
            print('Add Failed sinInsArea:%s;Because of:%s'%(sinInsArea,msg))
            driver.switch_to_alert().accept()
def addInsArea_Post(requestInfo,sinInsArea):
    if sinInsArea is None or len(sinInsArea) == 0:
        print('ID error,add ignore!')
    else:
        needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/InsureSystemSetEdit.aspx'
        needPostData = {}
        needPostData['__VIEWSTATE'] = get_ViewState(needurl,requestInfo.headers,requestInfo.http)
        needPostData['txtAreaName'] = sinInsArea
        needPostData['btnAdd'] = '新增'
        needPostData['__VIEWSTATEGENERATOR'] = '97C38CD7'
        response = requestNeedurl(needurl,postData = needPostData,headers = requestInfo.headers,CookieOpener = requestInfo.CookieOpener)
        readSoup = bs.BeautifulSoup(response.read().decode(),'html.parser')   
        msg = re.findall(r"alert\(\'.*\'\)",str(readSoup))
        if msg is None or len(msg) == 0:
            print('No Msg!')
        else:
            print('Add Msg:%s'%(msg[0])) 
def queryInsArea(requestInfo):
    listInsArea = []
    needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/InsureSystemSet.aspx'    
    content = requestInfo.http.request(uri = needurl,method = 'GET', headers=requestInfo.headers)[1]
    readSoup = bs.BeautifulSoup(content.decode(),'lxml')
    listInsArea = listInsArea + [singleGroup.find_all('span')[2].string for singleGroup in readSoup.find('div',id='M_TreeInsureArea_1').find_all('div') if len(singleGroup.find_all('span')[2].string)>0].copy()     
    print('listInsArea:%s'%(listInsArea))
    return   listInsArea  
def updateInsAreaByDoc(driver,requestInfo,listPublicByDoc):
    tmplist_InsArea_ByDoc = [sinPC.get('codeItem') for sinPC in listPublicByDoc if sinPC.get('codeName') == '缴纳地区']
    list_InsArea_ByDoc = [sinCI.get('codeItemValue') for sinCI in tmplist_InsArea_ByDoc[0]]
    list_InsArea_Sys = queryInsArea(requestInfo)
    for sin_InsArea_ByDoc in list_InsArea_ByDoc:
        if sin_InsArea_ByDoc in list_InsArea_Sys:
            print('sin_InsArea_ByDoc %s in list_InsArea_Sys,Add ignore!'%(sin_InsArea_ByDoc))
        elif len(sin_InsArea_ByDoc) > 0:
            addInsArea(driver,requestInfo,sin_InsArea_ByDoc)
    print('updateInsAreaByDoc End:%s'%(list_InsArea_ByDoc))
def queryInsAreaID(driver,requestInfo):
    '''request web,and click insarea one by one fro every InsAreaID
        return dic for example:{'无锡': '63940cc8-0842-47d5-9102-1bff4c56dd53', '特殊': 'f950e623-6380-4568-86b4-8acb20fc7127'}
    '''
    listInsArea = queryInsArea(requestInfo)
    needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/InsureSystemSet.aspx' 
    dictInsArea = {}
    driver.get(needurl)
    time.sleep(2)
    for i in range(len(listInsArea)):#range(len(listInsArea))
        driver.implicitly_wait(2)
        driver.find_element_by_xpath("//div[@id='TreeInsureArea_1_"+ str(i + 1) +"']/span[3]").click()
        driver.implicitly_wait(2)
        tmpPageSource = driver.page_source    
        readSoup = bs.BeautifulSoup(tmpPageSource,"html.parser")
        try:
            systemid =dict(readSoup.findAll(id = 'txtSystemID')[0].attrs)['value'] 
        except(exception):
            print('readSoup:%s'%(readSoup))
        else:
            print('listInsArea:%s;systemid:%s'%(listInsArea[i],systemid))
        dictInsArea[listInsArea[i]] = systemid       
    print('dictInsArea:%s'%(dictInsArea))
    return dictInsArea
    '''
    needurl = requestInfo.rooturl + r'/ePayroll/PersonalPayrollInformationManage/PersonalPayrollStructureGroupInput.aspx?Module=INS'
    content = requestInfo.http.request(uri = needurl,method = 'GET', headers=requestInfo.headers)[1]
    readSoup = bs.BeautifulSoup(content.decode(),'lxml')
    listvalue = {sinoption.string:sinoption.get('value') for sinoption in readSoup.find('select',id='UcStaffSelect1_ddlqarea').find_all('option') if sinoption.string is not None}    
    print('listvalue:%s'%(listvalue))
    return  listvalue
    
    '''
def queryPersonPropID(requestInfo):
    needurl = requestInfo.rooturl + r'/ePayroll/PersonalPayrollInformationManage/PersonalPayrollStructureGroupInput.aspx?Module=INS'
    content = requestInfo.http.request(uri = needurl,method = 'GET', headers=requestInfo.headers)[1]
    readSoup = bs.BeautifulSoup(content.decode(),'lxml')
    listvalue = {sinoption.string:sinoption.get('value') for sinoption in readSoup.find('select',id='UcStaffSelect1_ddlqins').find_all('option') if sinoption.string is not None}    
    print('listvalue:%s'%(listvalue))
    return  listvalue
def queryInsTypeID(requestInfo):
    needurl = requestInfo.rooturl + r'/ePayroll/PersonalPayrollInformationManage/InsuranceDetailInfo.aspx'
    content = requestInfo.http.request(uri = needurl,method = 'GET', headers=requestInfo.headers)[1]
    readSoup = bs.BeautifulSoup(content.decode(),'lxml')
    listvalue = {sinoption.string:sinoption.get('value') for sinoption in readSoup.find('select',id='DdlInsuranceType').find_all('option') if sinoption.string is not None}    
    print('listvalue:%s'%(listvalue))
    return  listvalue
def addAreaDetail(driver,requestInfo,sinPerProID,sinInsAreaID,sinInsAreaDetail):
    if sinPerProID is None or sinInsAreaID is None or len(sinPerProID) == 0 or len(sinInsAreaID) == 0:
        print('ID error,add ignore!')
    else:
        needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/InsureSystemSetItemEdit.aspx?SystemID='+ sinInsAreaID + '&PersonPropID=' + sinPerProID
        driver.get(needurl)
        time.sleep(2)
        driver.refresh()
        time.sleep(2)
        kindName = sinInsAreaDetail.get('险种名称')
        orderNo = str(int(sinInsAreaDetail.get('排序')))
        keepMod = sinInsAreaDetail.get('取舍方式')
        keepBits = str(sinInsAreaDetail.get('保留小数位'))
        priPayMode = sinInsAreaDetail.get('个人缴费类型')
        priPayValue = str(sinInsAreaDetail.get('个人缴费金额或比例'))
        priPayLowLevel = str(sinInsAreaDetail.get('个人基数下限规则'))
        priPayUpLevel = str(sinInsAreaDetail.get('个人基数上限规则'))
        pubPayMode = sinInsAreaDetail.get('单位缴费类型')
        pubPayValue = str(sinInsAreaDetail.get('单位缴费金额或比例'))
        pubPayLowLevel = str(sinInsAreaDetail.get('单位基数下限规则'))
        pubPayUpLevel = str(sinInsAreaDetail.get('单位基数上限规则'))
        dateEffectDate_GuruDate = r'2017-01-01'
         
        driver.find_element_by_xpath("//select[@id='drpKindName']/option[@title='"+kindName+"']").click()
        driver.find_element_by_id('txtOrderNo').clear()
        driver.find_element_by_id('txtOrderNo').send_keys(orderNo)
        driver.find_element_by_xpath("//select[@id='drpKeepMod']/option[@title='"+keepMod+"']").click()
        driver.find_element_by_xpath("//select[@id='drpkeepBits']/option[@title='"+keepBits+"']").click()
        driver.find_element_by_id('txtDateEffectDate_GuruDate').send_keys(dateEffectDate_GuruDate)
        #if paymode is 金额,lowlevel and uplevel can't be write,so config it 比例 first,after write,config paymode last
        driver.find_element_by_xpath("//select[@id='drpPriPayMode']/option[@title='"+'比例'+"']").click()       
        driver.find_element_by_id('txtPriPayValue').clear()
        driver.find_element_by_id('txtPriPayValue').send_keys(priPayValue)
        driver.find_element_by_id('txtPriPayLowLevel').clear()
        driver.find_element_by_id('txtPriPayLowLevel').send_keys(priPayLowLevel)
        driver.find_element_by_id('txtPriPayUpLevel').clear()
        driver.find_element_by_id('txtPriPayUpLevel').send_keys(priPayUpLevel)
        driver.find_element_by_xpath("//select[@id='drpPriPayMode']/option[@title='"+priPayMode+"']").click()
        driver.find_element_by_xpath("//select[@id='drpPubPayMode']/option[@title='"+'比例'+"']").click()
        driver.find_element_by_id('txtPubPayValue').clear()
        driver.find_element_by_id('txtPubPayValue').send_keys(pubPayValue)
        driver.find_element_by_id('txtPubPayLowLevel').clear()
        driver.find_element_by_id('txtPubPayLowLevel').send_keys(pubPayLowLevel)
        driver.find_element_by_id('txtPubPayUpLevel').clear()
        driver.find_element_by_id('txtPubPayUpLevel').send_keys(pubPayUpLevel)
        driver.find_element_by_xpath("//select[@id='drpPubPayMode']/option[@title='"+pubPayMode+"']").click()
        driver.find_element_by_id('btnAddInsureKind').click()
        time.sleep(2)
        try:
            msg = driver.switch_to_alert().text
        except(Exception):
            print('no msg for kindname:%s'%(kindName))
        else:
            if msg == '新增成功':
                print('Add Success kindname：%s'%(kindName))
                driver.switch_to_alert().accept()
            else:
                print('Add Failed kindname:%s;Because of:%s'%(kindName,msg))
                driver.switch_to_alert().accept()
        
        print('add kind detail:%s end;'%(kindName))         
        
def readInsAreaDetailByDoc(requestInfo,filename):
    #step1:read InsAreaDetailCfg from payrollcfgfile
    workbook = xlrd.open_workbook(filename)
    curws = workbook.sheet_by_name('附2保险缴纳体系')
    nrows,ncols = curws.usp_get_len()
    content_value = []
    for i in range(nrows):
        content_rowValue = []
        if i == 0:
            content_Name = [curws.cell_value(i,j) for j in range(ncols)]
        else:
            content_rowValue = [curws.cell_value(i,j) for j in range(ncols)]
            content_value.append(content_rowValue.copy())              
    #step2:gather it into dict    
    tableContent = [{content_Name[i]:sin_rowvalue[i] for i in range(len(sin_rowvalue))} for sin_rowvalue in content_value ]
    #print('tableContent:%s'%(tableContent))
    return tableContent
def updateInsAreaDetailByDoc(driver,requestInfo,list_InsAreaDetail_ByDoc):
    PC_InsType = queryInsTypeID(requestInfo)
    dictInsAreaID_Sys = queryInsAreaID(driver,requestInfo)
    dictPersonPropID_Sys = queryPersonPropID(requestInfo)
    print('请注意，由于需要添加个数为:%s;预估时间为:%s;'%(len(list_InsAreaDetail_ByDoc),len(list_InsAreaDetail_ByDoc)*5)) 
    for sin_InsAreaDetail_ByDoc in list_InsAreaDetail_ByDoc:
        sinPerPro_ByDoc = sin_InsAreaDetail_ByDoc.get('社保身份')
        sinInsArea_ByDoc = sin_InsAreaDetail_ByDoc.get('缴纳地区')
        sinKindName_ByDoc = sin_InsAreaDetail_ByDoc.get('险种名称')

        sinPerProID_Sys = dictPersonPropID_Sys.get(sinPerPro_ByDoc) 
        sinInsArea_Sys = dictInsAreaID_Sys.get(sinInsArea_ByDoc)
        print('缴纳地区:%s;社保身份:%s;险种:%s;'%(sinInsArea_ByDoc,sinPerPro_ByDoc,sinKindName_ByDoc))
        if sinPerProID_Sys is None or sinInsArea_Sys is None or len(sinPerProID_Sys) == 0 or len(sinInsArea_Sys) == 0:
            print('sinPerProID_Sys or sinInsArea_Sys  error,add ignore!Please update InsArea first')
        else:   
            addAreaDetail_Post(requestInfo,sinPerProID_Sys,sinInsArea_Sys,sin_InsAreaDetail_ByDoc,PC_InsType)
            #addAreaDetail(driver,requestInfo,sinPerProID_Sys,sinInsArea_Sys,sin_InsAreaDetail_ByDoc)
def get_ViewState(needurl,headers,http):
    content = http.request(uri = needurl,method = 'GET', headers=headers)[1]
    readSoup = bs.BeautifulSoup(content.decode(),'html.parser')   
    view_input = readSoup.find(id="__VIEWSTATE").get('value')      
    return view_input        
def addAreaDetail_Post(requestInfo,sinPerProID,sinInsAreaID,sinInsAreaDetail,PC_InsType):
    PC_KeepMod = {'四舍五入':'0','进位':'1','舍尾':'2','银行舍入法':'3','见分进元':'4','七舍八入(二舍三入)':'5'}
    PC_keepBits = {'0':'0','1':'1','2':'2','3':'3','4':'4'}
    PC_payMode = {'比例':'0','金额':'1'}
    if sinPerProID is None or sinInsAreaID is None or len(sinPerProID) == 0 or len(sinInsAreaID) == 0:
        print('ID error,add ignore!')
    else:
        needurl = requestInfo.rooturl + r'/ePayroll/PayrollParameterSetting/InsureSystemSetItemEdit.aspx?SystemID='+ sinInsAreaID + '&PersonPropID=' + sinPerProID     
        needPostData = {}
        needPostData['__VIEWSTATE'] = get_ViewState(needurl,requestInfo.headers,requestInfo.http)
        needPostData['drpKindName'] = PC_InsType.get(sinInsAreaDetail.get('险种名称'))
        needPostData['txtOrderNo'] = str(int(sinInsAreaDetail.get('排序')))
        needPostData['drpKeepMod'] = PC_KeepMod.get(sinInsAreaDetail.get('取舍方式'))
        needPostData['drpkeepBits'] = str(sinInsAreaDetail.get('保留小数位'))
        needPostData['drpPriPayMode'] = PC_payMode.get(sinInsAreaDetail.get('个人缴费类型'))
        needPostData['txtPriPayValue'] = str(sinInsAreaDetail.get('个人缴费金额或比例'))
        needPostData['txtPriPayLowLevel'] = str(sinInsAreaDetail.get('个人基数下限规则'))
        needPostData['txtPriPayUpLevel'] = str(sinInsAreaDetail.get('个人基数上限规则'))
        needPostData['drpPubPayMode'] = PC_payMode.get(sinInsAreaDetail.get('单位缴费类型'))
        needPostData['txtPubPayValue'] = str(sinInsAreaDetail.get('单位缴费金额或比例'))
        needPostData['txtPubPayLowLevel'] = str(sinInsAreaDetail.get('单位基数下限规则'))
        needPostData['txtPubPayUpLevel'] = str(sinInsAreaDetail.get('单位基数上限规则'))
        needPostData['txtDateEffectDate$GuruDate'] = '2017-01-01'  #txtDateEffectDate$GuruDate
        needPostData['btnAddInsureKind'] = '新增' #新增
        needPostData['__VIEWSTATEGENERATOR'] = '97C38CD7'
        #response, content = requestInfo.http.request(needurl,method = 'POST', headers=requestInfo.headers, body=urllib.parse.urlencode(needPostData).encode('utf-8'))
        #readSoup = bs.BeautifulSoup(content.decode(),'html.parser')   
        #msg = re.findall(r"alert\(.*\)",str(readSoup))
        
        response = requestNeedurl(needurl,postData = needPostData,headers = requestInfo.headers,CookieOpener = requestInfo.CookieOpener)
        readSoup = bs.BeautifulSoup(response.read().decode(),'html.parser')   
        msg = re.findall(r"alert\(.*\)",str(readSoup))
        if msg is None or len(msg) == 0:
            print('No Msg!')
        else:
            print('Add Msg:%s'%(msg[0]))
        
      

def projectend():
    pass   


cookie = cookielib.CookieJar()
handler = urllib.request.HTTPCookieProcessor(cookie)
CookieOpener = urllib.request.build_opener(handler)
headers = {}
headers['Accept'] = 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8'
headers['Accept-Encoding'] = 'deflate'
headers['Accept-Language'] = 'zh-CN,zh;q=0.8,en-US;q=0.5,en;q=0.3'
headers['Connection'] = 'keep-alive'
headers['User-Agent'] = 'Mozilla/5.0 (Windows NT 10.0; WOW64; rv:51.0) Gecko/20100101 Firefox/51.0'

rooturl = r'http://peter/zybxehr'
host = r'peter'
port = r'80'
webname = r'sa'
webpsw = r'111111'
tmpFilePath = r'E:\KM\GITPROJECT\HumanResourcesAss\AssForGaia\src\Guru\tmpfile'
showMode = r'show'
loginurl = rooturl + r'/Account/Logon?'
http = httplib2.Http()
requestInfo = UrlRequest(rooturl=rooturl,host=host,port=port,webname=webname,webpsw=webpsw,tmpFilePath=tmpFilePath,showMode=showMode,cookie=cookie,headers=headers,loginurl=loginurl,CookieOpener=CookieOpener,http=http)


resLogin = loginInSystem(requestInfo)
driver = createDriver(requestInfo)
#driver = ''
#事项1：通过doc读取公用代码
fileName = r'E:\工作文档\1608-中银\06业务流程分析\薪资管理\test.docx'
listPublicByDoc = readPublicByDoc(requestInfo,fileName)
listExtendByDoc = readExtendByDoc(requestInfo,fileName)

#print('list_InsArea_ByDoc:%s'%(list_InsArea_ByDoc))
#事项2：读取已存在的保险公用代码及薪资公用代码 更新公用代码群组时会自动调用
#listPublicGroupName = queryPublicGroup(requestInfo)
#事项3：新增公用代码群组 更新公用代码群组时会自动调用
#resAddPublicGroup = addPublicGroup(driver,requestInfo,'薪资公用代码','test002')
#事项4：更新公用代码群组
#updatePublicGroupByDoc(driver,requestInfo,listPublicByDoc)
#addSingleExtendGroup = {'codeItem': [{'列字段名称': '人员险种', '数据关联': '人员险种(人事同步)', '是否必填': '是', '数据类型': '文本', '显示方式': '选择框', '序号': '1'}, {'列字段名称': '财务代码', '数据关联': '', '是否必填': '是', '数据类型': '文本', '显示方式': '文本框', '序号': '2'}], 'extendDesc': '为费用凭证报表产品代码列维护员工险种编码与财务码对应关系，维护在根组织（0总公司）上。', 'extendType': '组织业务数据', 'extendName': '员工险种-财务码', 'repeatLimit': '无限制'}
#addSingleExtendDetail = addSingleExtendGroup.get('codeItem')[0]
#addExtendGroup(driver,requestInfo,addSingleExtendGroup,listExtendGroupName)
#事项5：查询业务数据群组
#listExtendGroup = queryExtendGroup(requestInfo)
#事项6：查询业务数据群组明细 根据业务数据群组ID
#listExtendDetail = queryExtendDetail(requestInfo,'90184dd3-25ba-4475-a15a-ab5e0811071d')
#事项7：添加业务数据明细 向指定的业务数据群组ID
#addExtendDetail(driver,requestInfo,addSingleExtendDetail,listExtendDetail,'90184dd3-25ba-4475-a15a-ab5e0811071d')
#事项5：更新公用代码明细
#importPublicCodeDetail(driver,requestInfo,listExtendByDoc,listPublicByDoc)
#事项6：更新业务数据群组 根据文档
#updateExtendGroup(driver,requestInfo,listExtendByDoc)
#事项7：更新辅助函数
#updateFormula(driver,requestInfo)
#事项8：更新自定义程序池
#updatePool(driver,requestInfo)
#事项：新增缴纳地区
sinInsArea = '测试'
addInsArea_Post(requestInfo,sinInsArea)
#queryInsArea(requestInfo)
queryInsAreaID(driver,requestInfo)
#queryPersonPropID(requestInfo)
#sinInsAreaID = r'b2130c41-da2b-4dcb-b6fe-ec2bed07dfb5'
#sinPerProID = r'952dca26-3736-4233-ab19-04ea8e7c6c50'
#payrollcfg_fileName = r'E:\工作文档\1608-中银\06业务流程分析\薪资管理\test.xlsx'
#list_InsAreaDetail_ByDoc = readInsAreaDetailByDoc(requestInfo,payrollcfg_fileName)

#sinInsAreaDetail = {'取舍方式': '四舍五入', '险种名称': '养老保险', '保留小数位': '2', '社保界定日默认维护在1号': '1', '社保身份排序': 2.0, '生效日期': '2017-01-01', '单位缴费类型': '比例', '单位基数下限规则': '2849', '缴纳地区': '北京', '单位缴费金额或比例': 0.19, '个人基数上限规则': 14244.0, '单位基数上限规则': 14244.0, '社保身份': '本地城镇', '个人基数下限规则': '2849', '系统标签': '系统标签', '排序': 1.0, '个人缴费金额或比例': 0.08, '个人缴费类型': '比例'}   
    
#addAreaDetail(driver,requestInfo,sinPerProID,sinInsAreaID,sinInsAreaDetail)
#addAreaDetail_Post(requestInfo,sinPerProID,sinInsAreaID,sinInsAreaDetail)

#updateInsAreaDetailByDoc(driver,requestInfo,list_InsAreaDetail_ByDoc)
t2 = time.clock()
print ("the project costs %.9fs"%(t2-t1))
