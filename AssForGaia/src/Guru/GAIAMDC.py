#encoding:utf8
import sys
import tkinter.messagebox as mb #用于在窗口中创建messagebox
import os #os.urandom用于随机生成秘钥key
import pymssql #用于连接mssql数据库
from argcomplete.compat import str
from twisted.conch.test.test_helper import HEIGHT
from openpyxl.styles.borders import Side
import codecs #用于查看编码格式
from PRPCRYPT import prpcrypt #调用加密算法，用于加密数据库链接中的密码
import json #json数据格式，用于写入数据库链接配置信息
import random
import string
from odo.backends.tests.test_json import json_file
from tkinter.messagebox import showinfo
from _codecs import encode
import urllib.request
import http.cookiejar as cookielib
import urllib.parse
from html.parser import HTMLParser
import html
import cchardet as chardet#获取指定页面的编码方式
from _codecs import decode
import gzip
import bs4 as bs
from Guru.main import writeToTxt,readFromTxt,requestNeedurl,requestNeedurlByJsonPost
import re
import math
import docx
from docutils.nodes import paragraph
import copy
import http.client as httplib
import httplib2
#addNewTmpl,queryTmpl,addTmplVar,dictQueryTmplVars,addNewTask,queryTaskbyTmpl,addMainDataSource,addSubDataSource,dictQuerySubDataSourceList,readDocument
#,listVariableMappingByTaskId,listRecipientVarsByTmpId,addVarMapping,addRecipient,addTmplContent
   
def addNewTmpl(name,requestInfo):
    '''Add new tmpl in project which projectId = 2ae753f5-7402-4fa4-88de-2df90a0024ea,you need apply name'''
    print('11114555:%s'%('11114555'))
    restest = queryTmpl(requestInfo=requestInfo)
    tmplId = [singleTmpl['id'] for singleTmpl in queryTmpl(requestInfo=requestInfo) if singleTmpl['name'] == name]
    if len(tmplId)==0:    
        needurl = requestInfo.rooturl + r'/mdc/Tmpl/ADD'
        print('needurl:%s'%(needurl))
        needPostData = {}
        needPostData['projectId'] = '2ae753f5-7402-4fa4-88de-2df90a0024ea' #默认职位人事业务
        needPostData['messageTypeId'] = '059dda48-f066-4ffe-b47b-69655434d90a'
        needPostData['priority'] = '1'
        needPostData['name'] = name
        needPostData['id'] = ''
        response = requestNeedurl(needurl=needurl,postData=needPostData,headers = requestInfo.headers,CookieOpener = requestInfo.CookieOpener)
        return response
    else:
        print('%s exist tmpl,skip Add'%(name))
def queryTmpl(requestInfo):
    '''Query tmpl list in project which projectId = 2ae753f5-7402-4fa4-88de-2df90a0024ea '''
    needurl = requestInfo.rooturl + r'/mdc/Tmpl/QueryTmpl'
    needPostData = {}
    needPostData['projectId'] = '2ae753f5-7402-4fa4-88de-2df90a0024ea' #默认职位人事业务
    needPostData['pageIndex'] = '0'
    needPostData['pageSize'] = '30'
    response = requestNeedurl(needurl=needurl,postData=needPostData,headers = requestInfo.headers,CookieOpener = requestInfo.CookieOpener)
    resQueryTmpl = json.loads(response.read().decode())
    josonQueryTmpl = resQueryTmpl['data']['records']
    return josonQueryTmpl
def addTmplVar(templateId,name,code,requestInfo):
    '''Add TmplVar to Tmpl which templateId = tmplId,you need apply name and code,def will check exists'''
    needurl = requestInfo.rooturl + r'/mdc/Tmpl/AddVariable'
    needPostData = {}
    needPostData['id'] = ''
    needPostData['templateId'] = templateId
    needPostData['name'] = name
    needPostData['code'] = code
    dictTmplVar = dictQueryTmplVars(tmplId=templateId,requestInfo=requestInfo)
    listTmpVarName = []
    for singleTmpVar in dictTmplVar:
        listTmpVarName.append(singleTmpVar['name'])
    #print(listTmpVarName)
    if name not in listTmpVarName:
        response = requestNeedurl(needurl=needurl,postData=needPostData,headers = requestInfo.headers,CookieOpener = requestInfo.CookieOpener) 
        print('%s Add Success'%(name))
    else:
        print('%s exist vars,skip Add'%(name))      
def dictQueryTmplVars(tmplId,requestInfo):
    '''Query all TmplVars IN Tmpl which tmplId = tmplId and resquest the dict of tmplvar list'''
    needurl = requestInfo.rooturl + r'/mdc/Tmpl/QueryVariable'
    needPostData = {}
    needPostData['tmplId'] = tmplId
    response = requestNeedurl(needurl=needurl,postData=needPostData,headers = requestInfo.headers,CookieOpener = requestInfo.CookieOpener)
    dictResponse = json.loads(response.read().decode())['data']
    return dictResponse
def addNewTask(name,tmplId,requestInfo):
    '''Add new task in tmpl,you need apply name and tmplId'''
    taskId = [singleTask['id'] for singleTask in queryTaskbyTmpl(tmplId=tmplId,requestInfo=requestInfo) if singleTask['name'] == name]
    if len(taskId) == 0:
        needurl = requestInfo.rooturl + r'/mdc/Task/ADD'
        needPostData = {}
        needPostData['templateId'] = tmplId #默认职位人事业务
        needPostData['name'] = name
        needPostData['id'] = ''
        needPostData['nameLangid'] = ''
        needPostData['isMail'] = 'true'
        needPostData['isMsg'] = 'false'
        needPostData['isSMS'] = 'false'
        needPostData['agentId'] = ''
        needPostData['isWeixin'] = 'false'
        needPostData['description'] = name
        needPostData['descriptionLangid'] = ''
        response = requestNeedurl(needurl=needurl,postData=needPostData,headers = requestInfo.headers,CookieOpener = requestInfo.CookieOpener)
        return response
    else:
        print('Add filed,taskname exists')
def queryTaskbyTmpl(tmplId,requestInfo):
    '''Query tmpl list in project which projectId = 2ae753f5-7402-4fa4-88de-2df90a0024ea '''
    needurl = requestInfo.rooturl + r'/mdc/Task/query'
    needPostData = {}
    needPostData['tmplId'] = tmplId
    needPostData['pageIndex'] = '0'
    needPostData['pageSize'] = '30'
    response = requestNeedurl(needurl=needurl,postData=needPostData,headers = requestInfo.headers,CookieOpener = requestInfo.CookieOpener)
    resQueryTask = json.loads(response.read().decode())
    josonQueryTask = resQueryTask['records']
    return josonQueryTask
def addMainDataSource(taskId,name,sqlStatement,requestInfo):
    '''Add MainDataSource to task which taskid = taskId,you need apply name and sqlStatement'''
    if taskId=='':
        print('Need taskId')
    else:
        needurl = requestInfo.rooturl + r'/mdc/Task/dataSource'
        needPostData = {}
        needPostData['id'] = ''
        needPostData['taskId'] = taskId
        needPostData['name'] = name
        needPostData['databaseSourceId'] = 'cb1ecd13-2dab-4892-b517-34d79ccfcb6d'#默认为ehr
        needPostData['nameLangid'] = ''
        needPostData['sqlStatement'] = sqlStatement
        response = requestNeedurl(needurl=needurl,postData=needPostData,headers = requestInfo.headers,CookieOpener = requestInfo.CookieOpener) 
        print('%s Add Success'%(name))
        return response
def addSubDataSource(taskId,name,sqlStatement,requestInfo):
    '''Add SubDataSource to task which taskId = taskId,you need apply name and sqlStatement,def will check exists'''
    #print(taskId)
    dictSubDataSource = dictQuerySubDataSourceList(taskId=taskId,requestInfo=requestInfo)
    #print(dictSubDataSource)
    listSubDataSourceName = []
    for singleSubDataSource in dictSubDataSource:
        listSubDataSourceName.append(singleSubDataSource['name'])
    #print(listSubDataSourceName)
    if name not in listSubDataSourceName:
        needurl = requestInfo.rooturl + r'/mdc/Task/AddDataSource'
        needPostData = {}
        needPostData['id'] = ''
        needPostData['taskId'] = taskId
        needPostData['name'] = name
        needPostData['databaseSourceId'] = 'cb1ecd13-2dab-4892-b517-34d79ccfcb6d'#默认为ehr
        needPostData['nameLangid'] = ''
        needPostData['sqlStatement'] = sqlStatement    
        response = requestNeedurl(needurl=needurl,postData=needPostData,headers = requestInfo.headers,CookieOpener = requestInfo.CookieOpener) 
        print('%s Add Success'%(name))
        return response
    else:
        print('%s exist SubDataSource,skip Add'%(name))  
def dictQuerySubDataSourceList(taskId,requestInfo):
    '''Query SubDataSourceList IN task which taskId = taskId and resquest the dict of SubDataSource list'''
    needurl = requestInfo.rooturl + r'/mdc/Task/SubDataSourceList'
    needPostData = {}
    needPostData['taskId'] = taskId
    needPostData['pageIndex'] = 0
    needPostData['pageSize'] = 30
    response = requestNeedurl(needurl=needurl,postData=needPostData,headers = requestInfo.headers,CookieOpener = requestInfo.CookieOpener)
    dictResponse = json.loads(response.read().decode())['records']
    return dictResponse
def readDocument(truepath):
    document = docx.Document(truepath)
    rowContent = []
    cellContent = []
    tableindex = 0
    paragraphindex = -1
    listParagraphData = [] 
         
    for table in  document.tables:
        tableindex += 1
        dictParagraphData = {} 
        if table.cell(0,0).text == '邮件提醒模板设定':
            paragraphindex += 1   
            rowContent = []
            rowDictContent = {}
            for row in table.rows:
                cellContent = []
                for cell in row.cells:
                    cellContent.append(copy.copy(cell.text))  
                if cellContent[0]=='':
                    pass
                else:
                    rowDictContent[cellContent[0]] = cellContent[1] 
            dictParagraphData[table.cell(0,0).text] = rowDictContent
            if len(listParagraphData)<(paragraphindex+1):
                listParagraphData.append(copy.copy(dictParagraphData))
            else:
                listParagraphData[paragraphindex].update(copy.copy(dictParagraphData)) 
            #print('result1_listParagraphData[paragraphindex]:%s'%(listParagraphData[paragraphindex]))   
        elif table.cell(0,0).text == '变量显示名称':
            rowContent = []
            rowDictContent = {}
            for row in table.rows:       
                cellContent = []                
                for cell in row.cells:
                    cellContent.append(copy.copy(cell.text))
                   
                rowDictContent['name'] = cellContent[0] 
                rowDictContent['code'] = cellContent[1]
                rowDictContent['valueType'] = cellContent[2]
                rowDictContent['value'] = cellContent[3]
                if rowDictContent['name']=='变量显示名称':
                    pass
                else:
                    rowContent.append(copy.copy(rowDictContent))
            dictParagraphData[table.cell(0,0).text] = rowContent
            if len(listParagraphData)<(paragraphindex+1):
                listParagraphData.append(copy.copy(dictParagraphData))
            else:
                #print('dictParagraphData:%s'%(copy.copy(dictParagraphData)))
                #print('before2_listParagraphData[paragraphindex]:%s'%(listParagraphData[paragraphindex]))   
                listParagraphData[paragraphindex].update(copy.copy(dictParagraphData))            
            #print('result2_listParagraphData[paragraphindex]:%s'%(listParagraphData[paragraphindex]))     
        elif table.cell(0,0).text == '接收人设定':
            rowDictContent = {}
            for row in table.rows:
                cellContent = []              
                for cell in row.cells:
                    cellContent.append(copy.copy(cell.text))
                cellContent[1] = re.split('[；;]', cellContent[1]).copy()    
                if cellContent[0] == '收件人':
                    rowDictContent['MailTo'] = cellContent[1]                  
                elif cellContent[0] == '抄送人':
                    rowDictContent['Cc'] = cellContent[1]
                elif cellContent[0] == '密送人':
                    rowDictContent['Bcc'] = cellContent[1] 
            dictParagraphData[table.cell(0,0).text] = rowDictContent
            if len(listParagraphData)<(paragraphindex+1) or len(listParagraphData) == 0:
                listParagraphData.append(copy.copy(dictParagraphData))
            else:         
                listParagraphData[paragraphindex].update(copy.copy(dictParagraphData))
            #print('result3_listParagraphData[paragraphindex]:%s'%(listParagraphData[paragraphindex]))                                  
            #tableMailContent.append(rowDictContent) 
        elif table.cell(0,0).text == '主数据源设定':
            rowDictContent = {}
            for row in table.rows:
                cellContent = []
                for cell in row.cells:
                    cellContent.append(copy.copy(cell.text))
                if cellContent[0] == '主数据源设定':
                    rowDictContent['dataSourceName']=cellContent[1]
                elif cellContent[0] != '':
                    rowDictContent['sqlStatement']=cellContent[0]
            dictParagraphData[table.cell(0,0).text] = copy.copy(rowDictContent)
            if len(listParagraphData)<(paragraphindex+1) or len(listParagraphData) == 0:
                listParagraphData.append(copy.copy(dictParagraphData))
            else:         
                listParagraphData[paragraphindex].update(copy.copy(dictParagraphData))        
        elif table.cell(0,0).text == '子数据源设定':
            rowContent = []
            rowDictContent = {}
            for row in table.rows:       
                cellContent = []                
                for cell in row.cells:
                    cellContent.append(copy.copy(cell.text))           
                rowDictContent['subDataSourceName'] = cellContent[0] 
                rowDictContent['subDataSourceSqlStatement'] = cellContent[1]
                if rowDictContent['subDataSourceName']=='子数据源设定':
                    pass
                else:
                    rowContent.append(copy.copy(rowDictContent))
            dictParagraphData[table.cell(0,0).text] = rowContent
            if len(listParagraphData)<(paragraphindex+1) or len(listParagraphData) == 0:
                listParagraphData.append(copy.copy(dictParagraphData))
            else:
                #print('dictParagraphData:%s'%(copy.copy(dictParagraphData)))
                #print('before2_listParagraphData[paragraphindex]:%s'%(listParagraphData[paragraphindex]))   
                listParagraphData[paragraphindex].update(copy.copy(dictParagraphData))            
            #print('result2_listParagraphData[paragraphindex]:%s'%(listParagraphData[paragraphindex]))
    return listParagraphData
def listVariableMappingByTaskId(taskId,requestInfo):
    needurl = requestInfo.rooturl + r'/mdc/Task/variableMapping?taskId='+taskId
    needPostData = {}
    response = requestNeedurl(needurl=needurl,postData=needPostData,headers = requestInfo.headers,CookieOpener = requestInfo.CookieOpener)
    readVariableMappings = response.read().decode() 
    listVariableMappings = re.findall(r"variableList':\[(.*)\],",str(readVariableMappings))
    if len(listVariableMappings) > 0:
        listVariableMappings = '['+listVariableMappings[0]+']'
        listVariableMappings = json.loads(listVariableMappings)
    return listVariableMappings
def listRecipientVarsByTmpId(tmplId,requestInfo):
    needurl = requestInfo.rooturl + r'/mdc/Tmpl/Recipient?tmplId='+tmplId
    needPostData = {}
    response = requestNeedurl(needurl=needurl,postData=needPostData,headers = requestInfo.headers,CookieOpener = requestInfo.CookieOpener)
    readQueryRecipientVars = response.read().decode() 
    listRecipientVars = re.findall(r'variableList:\[(.*)\],',str(readQueryRecipientVars))
    if len(listRecipientVars) > 0:
        listRecipientVars = '['+listRecipientVars[0]+']'
        listRecipientVars = json.loads(listRecipientVars)
    return listRecipientVars
def addVarMapping(taskId,cookie,requestInfo,singleParagraphData):
    if taskId=='':
        print('Need taskId')
    else:        
        needurl = requestInfo.rooturl + r'/mdc/Task/SaveVariableMapping'   
        needPostData = {}
        dictVarMappings = []
        listVarMappings = listVariableMappingByTaskId (taskId=taskId,requestInfo=requestInfo)      
        listSubDataSources = dictQuerySubDataSourceList(taskId=taskId,requestInfo=requestInfo)
        #print('listSubDataSources:%s'%(listSubDataSources)) 
        listAddNewVars = singleParagraphData.get('变量显示名称')
        for singleVarMapping in listVarMappings:
            dictVarMapping = {} 
            for singleAddNewVar in listAddNewVars:
                if singleVarMapping['code']==singleAddNewVar['code']:
                    dictMerged=singleVarMapping.copy()
                    dictMerged.update(singleAddNewVar)    
                    if dictMerged['valueType'] =='主数据源字段':
                        dictMerged['mappingValueType'] = 3
                        dictMerged['mappingValue'] = dictMerged['value']
                    elif dictMerged['valueType'] =='子数据源':
                        for singleSubDataSource in listSubDataSources:
                            if singleSubDataSource['name']==singleAddNewVar['code']:
                                dictMerged['mappingValue'] = singleSubDataSource['id']
                        dictMerged['mappingValueType'] = 4
                    elif dictMerged['valueType'] =='系统参数':
                        dictMerged['mappingValueType'] = 2
                        if dictMerged['value'].lower() =='yyyy-MM-dd'.lower():
                            dictMerged['mappingValue'] = '1' 
                        elif  dictMerged['value'].lower() =='yyyy-MM-dd HH:mm:ss'.lower():
                            dictMerged['mappingValue'] = '2'
                        elif  dictMerged['value'].lower() =='yyyy-MM-dd HH:mm'.lower():
                            dictMerged['mappingValue'] = '3' 
                        elif  dictMerged['value'].lower() =='HH:mm:ss'.lower():
                            dictMerged['mappingValue'] = '4'    
                        elif  dictMerged['value'].lower() =='HH:mm'.lower():
                            dictMerged['mappingValue'] = '5'                
                    elif dictMerged['valueType'] =='常量':
                        dictMerged['mappingValueType'] = dictMerged['value'] 
            dictVarMapping['variableId'] = dictMerged['id']
            dictVarMapping['mappingValueType'] = dictMerged['mappingValueType']
            dictVarMapping['mappingValue'] = dictMerged['mappingValue']
            dictVarMappings.append(copy.copy(dictVarMapping))
        needPostData['taskId'] = taskId
        needPostData['models'] = dictVarMappings
        response = requestNeedurlByJsonPost(needurl=needurl,postData=needPostData,headers=requestInfo.headers,cookie=cookie,host=requestInfo.host,port=requestInfo.port)
        readMessage = json.loads(response.read().decode())
        print('VarMapping Update result:%s'%(readMessage['message']))        
        return response
def addRecipient(tmplId,cookie,listDocRecipientVars,requestInfo):
    if tmplId=='':
        print('Need tmplId')
    else:        
        needurl = requestInfo.rooturl + r'/mdc/Tmpl/SaveRecipient'   
        needPostData = {}
        dictRecipientVars = {}
        singleVar = {}
        RecipientVars = []
        listRecipientVars = listRecipientVarsByTmpId(tmplId=tmplId,requestInfo=requestInfo)
        #print('listRecipientVars:%s'%(listRecipientVars))
        #listAddNewVars = singleParagraphData.get('变量显示名称')
        for singleRecipientVar in listRecipientVars:
            dictRecipientVars[singleRecipientVar.get('name')] = singleRecipientVar.get('id')
        #print('dictRecipientVars:%s'%(dictRecipientVars)) 
        for key,value in listDocRecipientVars.items():
            i = 0
            sendType = 0            
            for item in value:
                if len(item) > 0:
                    #print('item:%s value:%s'%(key,item))
                    i += 1
                    if key == 'MailTo':
                        sendType = 1
                    elif key == 'Cc':
                        sendType = 2
                    elif key == 'Bcc':
                        sendType = 3      
                    objectId = dictRecipientVars.get(item)
                    singleVar['objectId'] = objectId
                    singleVar['objectType'] = 1
                    singleVar['sendType'] = sendType
                    singleVar['serialNo'] = i
                    RecipientVars.append(singleVar.copy())
        #print('RecipientVars:%s'%(RecipientVars))           
        needPostData['tmplId'] = tmplId
        needPostData['model'] = RecipientVars.copy()
        needPostData['scopeType'] = 1     
        #print('needPostData:%s'%(needPostData)) 
        response = requestNeedurlByJsonPost(needurl=needurl,postData=needPostData,headers=requestInfo.headers,cookie=cookie,host=requestInfo.host,port=requestInfo.port)
        readMessage = json.loads(response.read().decode())
        print('Recipient Update result:%s'%(readMessage['message']))        
        return response 
def addTmplContent(tmplId,cookie,designSubject,designContents,requestInfo):
    if tmplId=='':
        print('Need taskId') 
    else: 
        needurl = requestInfo.rooturl + r'/file/group'
        needPostData = {}    
        needPostData['id'] = tmplId
        response = requestNeedurlByJsonPost(needurl=needurl,postData=needPostData,headers=requestInfo.headers,cookie=requestInfo.cookie,host=requestInfo.host,port=requestInfo.port)
        del needPostData    
        readMessage = json.loads(response.read().decode()) 
        print(readMessage)  
        needurl = requestInfo.rooturl + r'/mdc/Tmpl/Publish'   
        needPostData = {}    
        needPostData['templateId'] = tmplId
        needPostData['id'] = ''
        needPostData['templateType'] = 1 
        needPostData['designSubject'] = designSubject 
        needPostData['designContents'] = designContents 
        #needPostData['fileGroupId'] = tmplId 
        print('needPostData:%s'%(needPostData)) 
         
        response = requestNeedurlByJsonPost(needurl=needurl,postData=needPostData,headers=requestInfo.headers,cookie=requestInfo.cookie,host=requestInfo.host,port=requestInfo.port)
        readMessage = json.loads(response.read().decode())
        print('Add tmpl content result:%s'%(readMessage['message']))        
        return response      
def endOfDef():
    pass